import os
import csv
import subprocess
import sys
import tempfile
import typer
import requests
import json
import logging
import re as _re
import shutil
import time
import stat
from typing import Optional, List
from pathlib import Path
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, DownloadColumn, TransferSpeedColumn
from dotenv import load_dotenv, find_dotenv
from shutil import which
from litellm import completion

# --- Optional bundled conversion libraries ---
# These provide Python-native file conversion without external tools.
try:
    from PIL import Image as PILImage
    _HAS_PILLOW = True
except ImportError:
    _HAS_PILLOW = False

try:
    import pillow_heif
    if _HAS_PILLOW:
        pillow_heif.register_heif_opener()
    _HAS_PILLOW_HEIF = True
except ImportError:
    _HAS_PILLOW_HEIF = False

try:
    from pydub import AudioSegment
    _HAS_PYDUB = True
except ImportError:
    _HAS_PYDUB = False

try:
    from docx import Document as DocxDocument
    _HAS_PYTHON_DOCX = True
except ImportError:
    _HAS_PYTHON_DOCX = False

try:
    import markdown as _markdown_lib
    _HAS_MARKDOWN = True
except ImportError:
    _HAS_MARKDOWN = False

try:
    from fpdf import FPDF
    _HAS_FPDF = True
except ImportError:
    _HAS_FPDF = False

try:
    from pypdf import PdfReader, PdfWriter
    _HAS_PYPDF = True
except ImportError:
    _HAS_PYPDF = False

try:
    import yaml as _yaml_lib
    _HAS_PYYAML = True
except ImportError:
    _HAS_PYYAML = False

try:
    import openpyxl
    _HAS_OPENPYXL = True
except ImportError:
    _HAS_OPENPYXL = False

try:
    import toml as _toml_lib
    _HAS_TOML = True
except ImportError:
    _HAS_TOML = False

try:
    import ebooklib
    from ebooklib import epub
    _HAS_EBOOKLIB = True
except ImportError:
    _HAS_EBOOKLIB = False

# --- GLOBAL CONFIG (configurable for testing) ---
_CONFIG = {
    "LOG_FILE": os.path.expanduser("~/Desktop/alfred_debug.log"),
    "AI_PROVIDER": "ollama",  # Options: ollama, openai, anthropic, google, etc.
    "AI_MODEL": "qwen3:4b",  # Model name (format depends on provider)
    "OLLAMA_API_BASE": "http://localhost:11434",  # For Ollama only
    "TEMPERATURE": 0.2,
    "APP_SUPPORT_DIR": Path.home() / "Library/Application Support/Alfred",
}

def _init_config():
    """Initialize configuration from environment. Call this at app startup."""
    # Load environment variables
    dotenv_path = find_dotenv(usecwd=True)
    if dotenv_path:
        load_dotenv(dotenv_path)
    
    # Override config from env
    _CONFIG["AI_PROVIDER"] = os.getenv("AI_PROVIDER", _CONFIG["AI_PROVIDER"])
    _CONFIG["AI_MODEL"] = os.getenv("AI_MODEL", _CONFIG["AI_MODEL"])
    _CONFIG["OLLAMA_API_BASE"] = os.getenv("OLLAMA_API_BASE", _CONFIG["OLLAMA_API_BASE"])
    _CONFIG["TEMPERATURE"] = float(os.getenv("TEMPERATURE", str(_CONFIG["TEMPERATURE"])))
    
    # Set up logging
    logging.basicConfig(
        filename=_CONFIG["LOG_FILE"],
        level=logging.DEBUG,
        format="%(asctime)s - %(levelname)s - %(message)s"
    )
    
    # Create tool directories
    local_bin_dir = _CONFIG["APP_SUPPORT_DIR"] / "bin"
    local_bin_dir.mkdir(parents=True, exist_ok=True)
    
    # Add local bin to PATH
    os.environ["PATH"] = f"{local_bin_dir}:{os.environ.get('PATH', '')}"

# Accessors for config values
def get_local_bin_dir() -> Path:
    return _CONFIG["APP_SUPPORT_DIR"] / "bin"

def get_ai_provider() -> str:
    return _CONFIG["AI_PROVIDER"]

def get_ai_model() -> str:
    return _CONFIG["AI_MODEL"]

def get_ollama_api_base() -> str:
    return _CONFIG["OLLAMA_API_BASE"]

def get_temperature() -> float:
    return _CONFIG["TEMPERATURE"]

console = Console()

TOOLS_URLS = {
    "ffmpeg": "https://evermeet.cx/ffmpeg/ffmpeg-7.1.zip",
    "pandoc": "https://github.com/jgm/pandoc/releases/download/3.6.3/pandoc-3.6.3-x86_64-macOS.zip",
}

# ============================================================
# UTILITIES
# ============================================================

def _human_size(nbytes: int) -> str:
    size = float(nbytes)
    for unit in ["B", "KB", "MB", "GB"]:
        if size < 1024:
            return f"{size:.0f} {unit}" if unit == "B" else f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"


def check_command_availability(cmd: str) -> bool:
    local_tool = get_local_bin_dir() / cmd
    if local_tool.exists() and os.access(local_tool, os.X_OK):
        return True
    return which(cmd) is not None


def get_llm_response(prompt: str, image_paths: Optional[List[str]] = None, retries: int = 2) -> str:
    """Get LLM response using LiteLLM (supports multiple providers and vision).
    
    Args:
        prompt: Text prompt for the model
        image_paths: Optional list of image file paths for vision models
        retries: Number of retry attempts
    
    Returns:
        str: Model response
    """
    import base64
    
    provider = get_ai_provider()
    model_name = get_ai_model()
    temperature = get_temperature()
    
    # Build full model identifier for LiteLLM
    if provider == "ollama":
        model = f"ollama/{model_name}"
        api_base = get_ollama_api_base()
    elif provider == "openai":
        model = f"openai/{model_name}"  # e.g., openai/gpt-4o
        api_base = None
    elif provider == "anthropic":
        model = f"anthropic/{model_name}"  # e.g., anthropic/claude-3-5-sonnet-20241022
        api_base = None
    elif provider == "google" or provider == "gemini":
        # Handle both "gemini-2.5-flash" and "gemini/gemini-2.5-flash" formats
        if not model_name.startswith("gemini/"):
            model = f"gemini/{model_name}"
        else:
            model = model_name
        api_base = None
    else:
        # For custom providers, assume format is already correct
        model = model_name
        api_base = None
    
    logging.info(f"Using {provider} provider with model {model}")
    
    for attempt in range(retries + 1):
        try:
            # Build message content
            if image_paths and len(image_paths) > 0:
                # Vision request with images
                content_parts = [{"type": "text", "text": prompt}]
                
                # Add images
                for img_path in image_paths[:5]:  # Limit to 5 images
                    if not os.path.exists(img_path):
                        logging.warning(f"Image not found: {img_path}")
                        continue
                    
                    # Read and encode image
                    with open(img_path, "rb") as f:
                        img_data = base64.b64encode(f.read()).decode("utf-8")
                    
                    # Detect image type
                    ext = Path(img_path).suffix.lower()
                    mime_map = {
                        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                        ".png": "image/png", ".gif": "image/gif",
                        ".webp": "image/webp", ".bmp": "image/bmp"
                    }
                    mime_type = mime_map.get(ext, "image/jpeg")
                    
                    content_parts.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{img_data}"
                        }
                    })
                
                message_content = content_parts
            else:
                # Text-only request
                message_content = prompt
            
            # Build completion request
            kwargs = {
                "model": model,
                "messages": [{"role": "user", "content": message_content}],
                "temperature": temperature,
                "timeout": 120
            }
            
            # Add api_base for Ollama
            if api_base:
                kwargs["api_base"] = api_base
            
            response = completion(**kwargs)
            content = response.choices[0].message.content.strip()
            
            # Strip <think> tags if present
            content = _re.sub(r"<think>.*?</think>", "", content, flags=_re.DOTALL).strip()
            logging.debug(f"LLM Response: {content[:300]}...")
            return content
            
        except Exception as e:
            error_msg = str(e).lower()
            
            # Check for connection errors
            if "connection" in error_msg or "connect" in error_msg:
                if attempt < retries:
                    console.print(f"[yellow]Connection failed, retrying ({attempt+1}/{retries+1})...[/yellow]")
                    time.sleep(2)
                    continue
                console.print(f"[bold red]Error:[/bold red] Cannot connect to {provider}. Check configuration.")
                return f"Error: Cannot connect to {provider}"
            
            # Check for timeout
            elif "timeout" in error_msg:
                if attempt < retries:
                    console.print(f"[yellow]Timeout, retrying ({attempt+1}/{retries+1})...[/yellow]")
                    continue
                return "Error: Request timed out"
            
            # Other errors
            else:
                logging.error(f"LLM Error: {e}", exc_info=True)
                if attempt < retries:
                    console.print(f"[yellow]Error occurred, retrying ({attempt+1}/{retries+1})...[/yellow]")
                    time.sleep(1)
                    continue
                return f"Error: {e}"
    
    return "Error: Failed after retries"


def extract_code_block(response: str) -> tuple:
    for lang in ["python", "bash", "sh"]:
        marker = f"```{lang}"
        if marker in response:
            try:
                code = response.split(marker, 1)[1].split("```", 1)[0].strip()
                if code:
                    return (lang if lang != "sh" else "bash", code)
            except IndexError:
                continue
    return (None, None)


DANGEROUS_PATTERNS = [
    "rm -rf /", "rm -rf ~", "mkfs", "dd if=", ":(){",
    "chmod -R 777 /", "> /dev/sda", "shutdown", "reboot",
]
DANGEROUS_REGEXES = [
    _re.compile(r"curl\s+.*\|\s*(sh|bash)", _re.IGNORECASE),
    _re.compile(r"wget\s+.*\|\s*(sh|bash)", _re.IGNORECASE),
]


def execute_shell_command(command: str) -> bool:
    """Execute a shell command with safety checks. Returns True on success, False on failure/block."""
    logging.info(f"Executing: {command}")
    cmd_lower = command.lower().strip()
    for p in DANGEROUS_PATTERNS:
        if p in cmd_lower:
            console.print(f"[bold red]Blocked:[/bold red] Dangerous command detected.")
            return False
    for r in DANGEROUS_REGEXES:
        if r.search(cmd_lower):
            console.print(f"[bold red]Blocked:[/bold red] Dangerous command detected.")
            return False

    console.print(f"[blue]$ {command}[/blue]")
    try:
        env = os.environ.copy()
        env["PATH"] = f"{get_local_bin_dir()}:{env.get('PATH', '')}"
        
        result = subprocess.run(
            command, 
            shell=True, 
            check=True, 
            text=True, 
            capture_output=True, 
            timeout=300,
            env=env
        )
        if result.stdout:
            console.print(result.stdout.rstrip())
        console.print("[green]Done.[/green]")
        return True
    except subprocess.TimeoutExpired:
        console.print("[bold red]Timed out (5 min limit).[/bold red]")
        return False
    except subprocess.CalledProcessError as e:
        logging.error(f"Command failed: {e.stderr}")
        console.print(f"[red]{e.stderr.rstrip()}[/red]")
        return False


def execute_python_script(script_content: str) -> bool:
    """Execute Python script in temp file. Returns True on success, False on failure."""
    logging.info("Executing Python script")
    logging.debug(f"Script:\n{script_content}")
    with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False) as tmp:
        tmp.write(script_content)
        tmp_path = tmp.name
    python_exe = which("python3") or which("python")
    if not python_exe:
        console.print("[bold red]Error:[/bold red] No Python interpreter found.")
        return False
    try:
        result = subprocess.run([python_exe, tmp_path], check=True, text=True, capture_output=True)
        if result.stdout:
            console.print(result.stdout.rstrip())
        console.print("[green]Done.[/green]")
        return True
    except subprocess.CalledProcessError as e:
        logging.error(f"Script failed: {e.stderr}")
        console.print(f"[red]{e.stderr.rstrip()}[/red]")
        return False
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


EXTENSION_CATEGORIES = {
    "Images": {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp", ".svg", ".ico", ".heic", ".heif"},
    "Documents": {".pdf", ".doc", ".docx", ".txt", ".rtf", ".odt", ".pages", ".tex", ".md", ".rst", ".epub"},
    "Spreadsheets": {".csv", ".xlsx", ".xls", ".tsv", ".ods", ".numbers"},
    "Audio": {".mp3", ".wav", ".flac", ".ogg", ".aac", ".m4a", ".wma", ".opus"},
    "Video": {".mp4", ".avi", ".mkv", ".mov", ".webm", ".flv", ".wmv", ".m4v"},
    "Archives": {".zip", ".tar", ".gz", ".bz2", ".rar", ".7z", ".xz", ".dmg", ".iso"},
    "Code": {".py", ".js", ".ts", ".html", ".css", ".java", ".c", ".cpp", ".h", ".swift", ".go", ".rs", ".rb", ".sh"},
    "Data": {".json", ".xml", ".yaml", ".yml", ".toml", ".sql", ".db", ".sqlite"},
    "Presentations": {".ppt", ".pptx", ".key", ".odp"},
    "Design": {".psd", ".ai", ".sketch", ".fig", ".xd"},
}


def _categorize_file(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    for category, extensions in EXTENSION_CATEGORIES.items():
        if ext in extensions:
            return category
    return "Other"


def _json_to_yaml_simple(obj, indent: int = 0) -> list:
    lines = []
    prefix = "  " * indent
    if isinstance(obj, dict):
        for k, v in obj.items():
            if isinstance(v, (dict, list)):
                lines.append(f"{prefix}{k}:")
                lines.extend(_json_to_yaml_simple(v, indent + 1))
            else:
                val = "true" if v is True else "false" if v is False else "null" if v is None else str(v)
                if isinstance(v, str) and (':' in v or '#' in v or '\n' in v):
                    val = f'"{v}"'
                lines.append(f"{prefix}{k}: {val}")
    elif isinstance(obj, list):
        for item in obj:
            if isinstance(item, (dict, list)):
                lines.append(f"{prefix}-")
                lines.extend(_json_to_yaml_simple(item, indent + 1))
            else:
                lines.append(f"{prefix}- {item}")
    else:
        lines.append(f"{prefix}{obj}")
    return lines


def _xml_to_dict(element) -> dict:
    result: dict = {}
    if element.attrib:
        result["@attributes"] = dict(element.attrib)
    for child in element:
        child_data = _xml_to_dict(child)
        tag = child.tag
        if tag in result:
            if not isinstance(result[tag], list):
                result[tag] = [result[tag]]
            result[tag].append(child_data)
        else:
            result[tag] = child_data
    if element.text and element.text.strip():
        if result:
            result["#text"] = element.text.strip()
        else:
            return element.text.strip()  # type: ignore[return-value]
    return result


def _convert_data(input_file: str, ext: str, target: str, output_path: str):
    """Deterministic data file conversion using Python stdlib."""
    key = f"{ext}->.{target}"
    logging.info(f"Data conversion: {key}")

    if ext == ".json" and target == "csv":
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if isinstance(data, dict): data = [data]
        if not isinstance(data, list) or len(data) == 0:
            console.print("[red]Error: JSON must be a list or dict for CSV conversion.[/red]")
            return False
        
        # Handle list of primitives (convert to single-column CSV)
        if not isinstance(data[0], dict):
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(["value"])  # Header
                for item in data:
                    writer.writerow([item])
            return True
        
        # Handle list of dicts (standard case)
        headers = list(data[0].keys())
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=headers)
            writer.writeheader()
            for row in data:
                writer.writerow(row)
        return True

    elif ext == ".csv" and target == "json":
        with open(input_file, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            data = list(reader)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True

    # --- Extended data conversions via bundled libraries ---

    # JSON <-> YAML
    if ext == ".json" and target in ("yaml", "yml") and _HAS_PYYAML:
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        with open(output_path, 'w', encoding='utf-8') as f:
            _yaml_lib.dump(data, f, default_flow_style=False, allow_unicode=True)
        return True

    if ext in (".yaml", ".yml") and target == "json" and _HAS_PYYAML:
        with open(input_file, 'r', encoding='utf-8') as f:
            data = _yaml_lib.safe_load(f)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True

    # JSON/CSV <-> XLSX
    if ext == ".json" and target == "xlsx" and _HAS_OPENPYXL:
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if isinstance(data, dict):
            data = [data]
        if not isinstance(data, list) or len(data) == 0:
            console.print("[red]Error: JSON must be a list or dict for XLSX conversion.[/red]")
            return False
        wb = openpyxl.Workbook()
        ws = wb.active
        if isinstance(data[0], dict):
            headers = list(data[0].keys())
            ws.append(headers)
            for row in data:
                ws.append([row.get(h, "") for h in headers])
        else:
            ws.append(["value"])
            for item in data:
                ws.append([item])
        wb.save(output_path)
        return True

    if ext == ".csv" and target == "xlsx" and _HAS_OPENPYXL:
        wb = openpyxl.Workbook()
        ws = wb.active
        with open(input_file, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                ws.append(row)
        wb.save(output_path)
        return True

    if ext == ".xlsx" and target == "csv" and _HAS_OPENPYXL:
        wb = openpyxl.load_workbook(input_file)
        ws = wb.active
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(row)
        return True

    if ext == ".xlsx" and target == "json" and _HAS_OPENPYXL:
        wb = openpyxl.load_workbook(input_file)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 1:
            json.dump([], open(output_path, 'w'), indent=2)
            return True
        headers = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
        data = []
        for row in rows[1:]:
            data.append({headers[i]: (v if v is not None else "") for i, v in enumerate(row)})
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False, default=str)
        return True

    # JSON <-> TOML
    if ext == ".json" and target == "toml" and _HAS_TOML:
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if not isinstance(data, dict):
            console.print("[red]Error: TOML requires a top-level dict/table.[/red]")
            return False
        with open(output_path, 'w', encoding='utf-8') as f:
            _toml_lib.dump(data, f)
        return True

    if ext == ".toml" and target == "json" and _HAS_TOML:
        with open(input_file, 'r', encoding='utf-8') as f:
            data = _toml_lib.load(f)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True

    console.print(f"[red]Error: No built-in converter for {ext} -> .{target}[/red]")
    return False


def _convert_with_pillow(input_file: str, target: str, output_path: str) -> bool:
    """Convert images using Pillow (bundled, no external tools needed)."""
    if not _HAS_PILLOW:
        console.print("[red]Error: Pillow is not installed.[/red]")
        return False
    try:
        img = PILImage.open(input_file)

        # Handle RGBA -> formats that don't support alpha
        if img.mode == "RGBA" and target.lower() in ("jpg", "jpeg", "bmp", "pdf"):
            background = PILImage.new("RGB", img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background
        elif img.mode == "P" and target.lower() in ("jpg", "jpeg"):
            img = img.convert("RGB")
        elif img.mode not in ("RGB", "L") and target.lower() in ("jpg", "jpeg"):
            img = img.convert("RGB")

        # Map target to Pillow format string
        pillow_format_map = {
            "jpg": "JPEG", "jpeg": "JPEG", "png": "PNG", "bmp": "BMP",
            "gif": "GIF", "tiff": "TIFF", "tif": "TIFF", "webp": "WEBP",
            "ico": "ICO", "pdf": "PDF",
        }
        fmt = pillow_format_map.get(target.lower())
        if not fmt:
            console.print(f"[red]Error: Pillow cannot write .{target}[/red]")
            return False

        # ICO requires specific sizes
        if fmt == "ICO":
            img = img.resize((256, 256), PILImage.LANCZOS)

        save_kwargs = {}
        if fmt == "JPEG":
            save_kwargs["quality"] = 90
        elif fmt == "WEBP":
            save_kwargs["quality"] = 85

        img.save(output_path, format=fmt, **save_kwargs)
        return True
    except Exception as e:
        logging.error(f"Pillow conversion failed: {e}", exc_info=True)
        console.print(f"[red]Error: Pillow conversion failed: {e}[/red]")
        return False


def _convert_with_pydub(input_file: str, target: str, output_path: str) -> bool:
    """Convert audio using pydub (bundled, uses ffmpeg if available, falls back to audioop)."""
    if not _HAS_PYDUB:
        console.print("[red]Error: pydub is not installed.[/red]")
        return False
    try:
        ext = Path(input_file).suffix.lower().lstrip(".")
        # Determine input format
        input_format_map = {
            "mp3": "mp3", "wav": "wav", "flac": "flac", "ogg": "ogg",
            "aac": "aac", "m4a": "m4a", "aiff": "aiff", "aif": "aiff",
        }
        in_fmt = input_format_map.get(ext, ext)
        audio = AudioSegment.from_file(input_file, format=in_fmt)

        # Determine output format
        output_format_map = {
            "mp3": "mp3", "wav": "wav", "flac": "flac", "ogg": "ogg",
            "aac": "adts", "m4a": "ipod", "aiff": "aiff",
        }
        out_fmt = output_format_map.get(target.lower(), target.lower())

        export_kwargs = {}
        if out_fmt == "mp3":
            export_kwargs["bitrate"] = "192k"
        elif out_fmt == "adts":
            export_kwargs["codec"] = "aac"
        elif out_fmt == "ipod":
            export_kwargs["codec"] = "aac"

        audio.export(output_path, format=out_fmt, **export_kwargs)
        return True
    except Exception as e:
        logging.error(f"pydub conversion failed: {e}", exc_info=True)
        console.print(f"[red]Error: pydub conversion failed: {e}[/red]")
        return False


def _convert_document_python(input_file: str, ext: str, target: str, output_path: str) -> bool:
    """Convert documents using bundled Python libraries."""
    try:
        # --- Markdown -> HTML ---
        if ext == ".md" and target == "html" and _HAS_MARKDOWN:
            with open(input_file, 'r', encoding='utf-8') as f:
                md_text = f.read()
            html = _markdown_lib.markdown(md_text, extensions=['tables', 'fenced_code', 'codehilite', 'toc'])
            full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{Path(input_file).stem}</title>
<style>body {{ font-family: -apple-system, sans-serif; max-width: 800px; margin: 2em auto; padding: 0 1em; line-height: 1.6; }}
pre {{ background: #f4f4f4; padding: 1em; overflow-x: auto; border-radius: 4px; }}
code {{ background: #f4f4f4; padding: 0.2em 0.4em; border-radius: 3px; }}
table {{ border-collapse: collapse; }} th, td {{ border: 1px solid #ddd; padding: 0.5em 1em; }}</style>
</head><body>{html}</body></html>"""
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            return True

        # --- Markdown -> PDF (via fpdf2) ---
        if ext == ".md" and target == "pdf" and _HAS_MARKDOWN and _HAS_FPDF:
            with open(input_file, 'r', encoding='utf-8') as f:
                md_text = f.read()
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Helvetica", size=11)
            for line in md_text.split('\n'):
                if line.startswith('### '):
                    pdf.set_font("Helvetica", style="B", size=13)
                    pdf.cell(0, 8, line[4:].strip(), new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", size=11)
                elif line.startswith('## '):
                    pdf.set_font("Helvetica", style="B", size=15)
                    pdf.cell(0, 10, line[3:].strip(), new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", size=11)
                elif line.startswith('# '):
                    pdf.set_font("Helvetica", style="B", size=18)
                    pdf.cell(0, 12, line[2:].strip(), new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", size=11)
                elif line.strip() == '':
                    pdf.ln(4)
                else:
                    pdf.multi_cell(0, 6, line)
            pdf.output(output_path)
            return True

        # --- HTML -> PDF (via fpdf2, text extraction) ---
        if ext == ".html" and target == "pdf" and _HAS_FPDF:
            from html.parser import HTMLParser

            class _HTMLTextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.lines: list[str] = []
                def handle_data(self, data):
                    text = data.strip()
                    if text:
                        self.lines.append(text)

            with open(input_file, 'r', encoding='utf-8') as f:
                html_content = f.read()
            parser = _HTMLTextExtractor()
            parser.feed(html_content)

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Helvetica", size=11)
            for line in parser.lines:
                pdf.multi_cell(0, 6, line)
                pdf.ln(2)
            pdf.output(output_path)
            return True

        # --- TXT -> HTML ---
        if ext == ".txt" and target == "html":
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()
            import html as _html_mod
            escaped = _html_mod.escape(text)
            paragraphs = escaped.split('\n\n')
            body = '\n'.join(f'<p>{p.replace(chr(10), "<br>")}</p>' for p in paragraphs if p.strip())
            full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{Path(input_file).stem}</title>
<style>body {{ font-family: -apple-system, sans-serif; max-width: 800px; margin: 2em auto; padding: 0 1em; line-height: 1.6; }}</style>
</head><body>{body}</body></html>"""
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            return True

        # --- TXT -> PDF (via fpdf2) ---
        if ext == ".txt" and target == "pdf" and _HAS_FPDF:
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Courier", size=10)
            for line in text.split('\n'):
                pdf.multi_cell(0, 5, line)
            pdf.output(output_path)
            return True

        # --- TXT/MD -> DOCX ---
        if ext in (".txt", ".md") and target == "docx" and _HAS_PYTHON_DOCX:
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()
            doc = DocxDocument()
            for line in text.split('\n'):
                # Basic markdown heading detection for .md files
                if ext == ".md" and line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                elif ext == ".md" and line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif ext == ".md" and line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                else:
                    doc.add_paragraph(line)
            doc.save(output_path)
            return True

        # --- DOCX -> TXT ---
        if ext == ".docx" and target == "txt" and _HAS_PYTHON_DOCX:
            doc = DocxDocument(input_file)
            text = '\n'.join(p.text for p in doc.paragraphs)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return True

        # --- DOCX -> PDF (via python-docx + fpdf2) ---
        if ext == ".docx" and target == "pdf" and _HAS_PYTHON_DOCX and _HAS_FPDF:
            doc = DocxDocument(input_file)
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Helvetica", size=11)
            for p in doc.paragraphs:
                if p.style and p.style.name.startswith('Heading'):
                    try:
                        level = int(p.style.name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    sizes = {1: 18, 2: 15, 3: 13, 4: 12}
                    pdf.set_font("Helvetica", style="B", size=sizes.get(level, 12))
                    pdf.cell(0, 10, p.text, new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", size=11)
                else:
                    if p.text.strip():
                        pdf.multi_cell(0, 6, p.text)
                    else:
                        pdf.ln(4)
            pdf.output(output_path)
            return True

        # --- HTML -> DOCX ---
        if ext == ".html" and target == "docx" and _HAS_PYTHON_DOCX:
            import html as _html_mod
            from html.parser import HTMLParser

            class _SimpleHTMLExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.parts: list[tuple[str, str]] = []
                    self._tag_stack: list[str] = []

                def handle_starttag(self, tag, attrs):
                    self._tag_stack.append(tag)

                def handle_endtag(self, tag):
                    if self._tag_stack and self._tag_stack[-1] == tag:
                        self._tag_stack.pop()

                def handle_data(self, data):
                    text = data.strip()
                    if not text:
                        return
                    current_tag = self._tag_stack[-1] if self._tag_stack else "p"
                    self.parts.append((current_tag, text))

            with open(input_file, 'r', encoding='utf-8') as f:
                html_content = f.read()
            parser = _SimpleHTMLExtractor()
            parser.feed(html_content)

            doc = DocxDocument()
            for tag, text in parser.parts:
                if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    level = int(tag[1])
                    doc.add_heading(text, level=level)
                elif tag in ("li",):
                    doc.add_paragraph(text, style="List Bullet")
                else:
                    doc.add_paragraph(text)
            doc.save(output_path)
            return True

        # --- HTML/MD -> EPUB ---
        if ext in (".html", ".md") and target == "epub" and _HAS_EBOOKLIB:
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Convert markdown to HTML first if needed
            if ext == ".md" and _HAS_MARKDOWN:
                content = _markdown_lib.markdown(content, extensions=['tables', 'fenced_code'])

            book = epub.EpubBook()
            title = Path(input_file).stem
            book.set_identifier(f'alfred-{title}')
            book.set_title(title)
            book.set_language('en')

            chapter = epub.EpubHtml(title=title, file_name='content.xhtml', lang='en')
            chapter.content = f'<html><body>{content}</body></html>'
            book.add_item(chapter)

            book.toc = [chapter]
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = ['nav', chapter]

            epub.write_epub(output_path, book)
            return True

        console.print(f"[red]Error: No Python-native converter for {ext} -> .{target}[/red]")
        return False

    except Exception as e:
        logging.error(f"Python document conversion failed: {e}", exc_info=True)
        console.print(f"[red]Error: Document conversion failed: {e}[/red]")
        return False


# Output format capabilities
SIPS_FORMATS = {"jpeg", "jpg", "png", "tiff", "tif", "bmp", "gif", "pict", "pdf", "heic"}
AFCONVERT_FORMATS = {"aac", "m4a", "wav", "aiff", "aif", "caf"}
TEXTUTIL_FORMATS = {"txt", "html", "rtf", "rtfd", "doc", "docx", "wordml", "odt", "webarchive"} # No PDF output
PANDOC_FORMATS = {"html", "pdf", "docx", "md", "rst", "tex", "epub", "txt", "rtf"}
FFMPEG_FORMATS = {"mp3", "wav", "aac", "m4a", "flac", "ogg", "mp4", "avi", "mkv", "mov", "webm", "gif"}
MAGICK_FORMATS = {"jpg", "jpeg", "png", "gif", "webp", "bmp", "tiff", "ico", "svg", "pdf"}

# --- Bundled Python library format capabilities ---
PILLOW_FORMATS = {"jpg", "jpeg", "png", "bmp", "gif", "tiff", "tif", "webp", "ico", "pdf", "heic", "heif"}
PYDUB_FORMATS = {"mp3", "wav", "flac", "ogg", "aac", "m4a", "aiff"}
PY_DOCX_FORMATS = {"docx"}  # python-docx can write .docx
PY_MARKDOWN_FORMATS = {"html"}  # markdown lib converts md -> html
PY_PDF_FORMATS = {"pdf"}  # fpdf2 generates PDFs (pure Python, no system deps)
PY_YAML_FORMATS = {"yaml", "yml"}  # PyYAML
PY_XLSX_FORMATS = {"xlsx"}  # openpyxl
PY_TOML_FORMATS = {"toml"}  # toml
PY_EPUB_FORMATS = {"epub"}  # ebooklib

def _tool_supports_target(tool: str, target: str) -> bool:
    """Check if a specific tool supports the target output format."""
    target = target.lower()
    if tool == "python": return True # Assumed specific logic handles it
    if tool == "sips": return target in SIPS_FORMATS
    if tool == "afconvert": return target in AFCONVERT_FORMATS
    if tool == "textutil": return target in TEXTUTIL_FORMATS
    if tool == "pandoc": return target in PANDOC_FORMATS
    if tool == "ffmpeg": return target in FFMPEG_FORMATS
    if tool == "magick": return target in MAGICK_FORMATS
    # Bundled Python libraries
    if tool == "pillow": return target in PILLOW_FORMATS
    if tool == "pydub": return target in PYDUB_FORMATS
    if tool == "py_docx": return target in PY_DOCX_FORMATS
    if tool == "py_markdown": return target in PY_MARKDOWN_FORMATS
    if tool == "py_pdf": return target in PY_PDF_FORMATS
    if tool == "py_yaml": return target in PY_YAML_FORMATS
    if tool == "py_xlsx": return target in PY_XLSX_FORMATS
    if tool == "py_toml": return target in PY_TOML_FORMATS
    if tool == "py_epub": return target in PY_EPUB_FORMATS
    return False

def _resolve_tool(tool_list: list[str]) -> str | None:
    for tool in tool_list:
        if tool == "python": return "python"
        elif tool == "sips": return "sips"
        elif tool == "afconvert": return "afconvert"
        elif tool == "textutil": return "textutil"
        elif tool == "ffmpeg" and check_command_availability("ffmpeg"): return "ffmpeg"
        elif tool == "magick" and (check_command_availability("magick") or check_command_availability("convert")): return "magick"
        elif tool == "pandoc" and check_command_availability("pandoc"): return "pandoc"
        # Bundled Python libraries (always available if installed)
        elif tool == "pillow" and _HAS_PILLOW: return "pillow"
        elif tool == "pydub" and _HAS_PYDUB: return "pydub"
        elif tool == "py_docx" and _HAS_PYTHON_DOCX: return "py_docx"
        elif tool == "py_markdown" and _HAS_MARKDOWN: return "py_markdown"
        elif tool == "py_pdf" and _HAS_FPDF: return "py_pdf"
        elif tool == "py_yaml" and _HAS_PYYAML: return "py_yaml"
        elif tool == "py_xlsx" and _HAS_OPENPYXL: return "py_xlsx"
        elif tool == "py_toml" and _HAS_TOML: return "py_toml"
        elif tool == "py_epub" and _HAS_EBOOKLIB: return "py_epub"
    return None


CONVERSION_MAP: dict[str, list[str]] = {
    # --- Data ---
    ".json->.csv": ["python"], ".csv->.json": ["python"],
    ".json->.yaml": ["py_yaml", "python"], ".json->.yml": ["py_yaml", "python"],
    ".yaml->.json": ["py_yaml", "python"], ".yml->.json": ["py_yaml", "python"],
    ".json->.xlsx": ["py_xlsx", "python"], ".csv->.xlsx": ["py_xlsx", "python"],
    ".xlsx->.csv": ["py_xlsx", "python"], ".xlsx->.json": ["py_xlsx", "python"],
    ".json->.toml": ["py_toml", "python"], ".toml->.json": ["py_toml", "python"],
    # --- Images (bundled Pillow as fallback) ---
    ".png->.jpg": ["sips", "magick", "pillow"], ".jpg->.png": ["sips", "magick", "pillow"],
    ".png->.webp": ["magick", "pillow"], ".webp->.png": ["magick", "sips", "pillow"],
    ".png->.bmp": ["sips", "magick", "pillow"], ".bmp->.png": ["sips", "magick", "pillow"],
    ".png->.gif": ["sips", "magick", "pillow"], ".gif->.png": ["sips", "magick", "pillow"],
    ".png->.tiff": ["sips", "magick", "pillow"], ".tiff->.png": ["sips", "magick", "pillow"],
    ".png->.ico": ["magick", "pillow"], ".jpg->.ico": ["magick", "pillow"],
    ".jpg->.webp": ["magick", "pillow"], ".webp->.jpg": ["magick", "sips", "pillow"],
    ".jpg->.bmp": ["sips", "magick", "pillow"], ".bmp->.jpg": ["sips", "magick", "pillow"],
    ".jpg->.gif": ["sips", "magick", "pillow"], ".gif->.jpg": ["sips", "magick", "pillow"],
    ".jpg->.tiff": ["sips", "magick", "pillow"], ".tiff->.jpg": ["sips", "magick", "pillow"],
    ".heic->.jpg": ["sips", "magick", "pillow"], ".heic->.png": ["sips", "magick", "pillow"],
    ".heif->.jpg": ["sips", "magick", "pillow"], ".heif->.png": ["sips", "magick", "pillow"],
    ".png->.pdf": ["sips", "magick", "pillow"], ".jpg->.pdf": ["sips", "magick", "pillow"],
    ".webp->.gif": ["magick", "pillow"], ".gif->.webp": ["magick", "pillow"],
    # --- Audio (bundled pydub as fallback) ---
    ".wav->.aac": ["afconvert", "ffmpeg", "pydub"], ".wav->.m4a": ["afconvert", "ffmpeg", "pydub"],
    ".mp3->.wav": ["ffmpeg", "afconvert", "pydub"], ".wav->.mp3": ["ffmpeg", "pydub"],
    ".mp3->.flac": ["ffmpeg", "pydub"], ".flac->.mp3": ["ffmpeg", "pydub"],
    ".mp3->.ogg": ["ffmpeg", "pydub"], ".ogg->.mp3": ["ffmpeg", "pydub"],
    ".wav->.flac": ["ffmpeg", "pydub"], ".flac->.wav": ["ffmpeg", "pydub"],
    ".wav->.ogg": ["ffmpeg", "pydub"], ".ogg->.wav": ["ffmpeg", "pydub"],
    ".m4a->.mp3": ["ffmpeg", "pydub"], ".m4a->.wav": ["ffmpeg", "pydub"],
    ".aac->.mp3": ["ffmpeg", "pydub"], ".aac->.wav": ["ffmpeg", "pydub"],
    # --- Video (ffmpeg only – no pure-Python fallback) ---
    ".mp4->.mp3": ["ffmpeg"], ".mp4->.wav": ["ffmpeg"],
    ".mp4->.avi": ["ffmpeg"], ".mp4->.mkv": ["ffmpeg"],
    ".mp4->.webm": ["ffmpeg"], ".mp4->.mov": ["ffmpeg"],
    ".avi->.mp4": ["ffmpeg"], ".mkv->.mp4": ["ffmpeg"],
    ".mov->.mp4": ["ffmpeg"], ".webm->.mp4": ["ffmpeg"],
    ".mp4->.gif": ["ffmpeg"],
    # --- Documents (bundled Python libs as fallback) ---
    ".txt->.html": ["textutil", "pandoc", "py_markdown"],
    ".txt->.docx": ["pandoc", "py_docx"],
    ".txt->.pdf": ["pandoc", "py_pdf"],
    ".docx->.pdf": ["pandoc", "py_pdf"],
    ".docx->.txt": ["pandoc", "py_docx"],
    ".md->.html": ["pandoc", "py_markdown"],
    ".md->.pdf": ["pandoc", "py_pdf"],
    ".md->.docx": ["pandoc", "py_docx"],
    ".html->.pdf": ["pandoc", "py_pdf"],
    ".html->.txt": ["textutil", "pandoc"],
    ".html->.docx": ["pandoc", "py_docx"],
    ".html->.epub": ["pandoc", "py_epub"],
    ".md->.epub": ["pandoc", "py_epub"],
}

SIPS_FORMAT_ALIAS = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "tiff": "tiff"}
AFCONVERT_FORMATS_MAP = {"aac": "aac ", "m4a": "m4af", "wav": "WAVE", "aiff": "AIFF"}



# ============================================================
# CLI APP
# ============================================================

app = typer.Typer(help="Alfred: Your Native Utility Agent")


@app.command()
def install(tool: str):
    """Download and install a tool locally (ffmpeg, pandoc)."""
    if tool not in TOOLS_URLS:
        console.print(f"[red]Error: Unknown tool '{tool}'.[/red]")
        console.print(f"[dim]Available: {', '.join(TOOLS_URLS.keys())}[/dim]")
        raise typer.Exit(1)
    
    url = TOOLS_URLS[tool]
    console.print(f"[blue]Downloading {tool}...[/blue]")
    
    zip_path = Path(tempfile.gettempdir()) / f"{tool}.zip"
    local_bin_dir = get_local_bin_dir()
    try:
        with requests.get(url, stream=True) as r:
            r.raise_for_status()
            total_size = int(r.headers.get('content-length', 0))
            
            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                DownloadColumn(),
                TransferSpeedColumn(),
                console=console
            ) as progress:
                task = progress.add_task(f"Fetching {tool}", total=total_size)
                with open(zip_path, 'wb') as f:
                    for chunk in r.iter_content(chunk_size=8192):
                        f.write(chunk)
                        progress.update(task, advance=len(chunk))
        
        console.print("[blue]Extracting...[/blue]")
        import zipfile
        with zipfile.ZipFile(zip_path, 'r') as z:
            bin_name = tool
            found = False
            for name in z.namelist():
                if name.endswith(f"/{tool}") or name == tool:
                    source = z.open(name)
                    target_path = local_bin_dir / tool
                    with open(target_path, 'wb') as target:
                        shutil.copyfileobj(source, target)
                    st = os.stat(target_path)
                    os.chmod(target_path, st.st_mode | stat.S_IEXEC)
                    found = True
                    break
            
            if not found:
                for name in z.namelist():
                    if f"bin/{tool}" in name:
                        source = z.open(name)
                        target_path = local_bin_dir / tool
                        with open(target_path, 'wb') as target:
                            shutil.copyfileobj(source, target)
                        st = os.stat(target_path)
                        os.chmod(target_path, st.st_mode | stat.S_IEXEC)
                        found = True
                        break

            if found:
                console.print(f"[green]Successfully installed {tool}![/green]")
            else:
                console.print(f"[red]Error: Could not find binary in zip archive.[/red]")
                
    except Exception as e:
        console.print(f"[red]Install failed: {e}[/red]")
    finally:
        if zip_path.exists():
            zip_path.unlink()


@app.command()
def convert(input_file: str, target_format: str):
    logging.info(f"Convert: {input_file} -> {target_format}")

    if not os.path.exists(input_file):
        console.print(f"[red]Error: File not found: {input_file}[/red]")
        raise typer.Exit(1)

    ext = os.path.splitext(input_file)[1].lower()
    # Normalize target: strip natural language like "convert to jpg", "to mp3", ".pdf"
    target = target_format.lower().strip()
    for prefix in ("convert to ", "convert to", "to ", "as ", "into "):
        if target.startswith(prefix):
            target = target[len(prefix):]
            break
    target = target.strip().lstrip(".").strip()
    output_path = os.path.splitext(input_file)[0] + f".{target}"
    key = f"{ext}->.{target}"

    tool_list = CONVERSION_MAP.get(key)

    if tool_list is None:
        # Heuristic guess (includes bundled Python libraries as fallbacks)
        all_media = EXTENSION_CATEGORIES.get("Audio", set()) | EXTENSION_CATEGORIES.get("Video", set())
        all_audio = EXTENSION_CATEGORIES.get("Audio", set())
        all_images = EXTENSION_CATEGORIES.get("Images", set())
        all_docs = EXTENSION_CATEGORIES.get("Documents", set())
        all_data = EXTENSION_CATEGORIES.get("Data", set()) | EXTENSION_CATEGORIES.get("Spreadsheets", set())

        if ext in all_data or f".{target}" in all_data:
            tool_list = ["python", "py_yaml", "py_xlsx", "py_toml"]
        elif ext in all_media or f".{target}" in all_media:
            # For audio-only, include pydub; video still needs ffmpeg
            if (ext in all_audio or f".{target}" in all_audio) and ext not in EXTENSION_CATEGORIES.get("Video", set()):
                tool_list = ["ffmpeg", "afconvert", "pydub"]
            else:
                tool_list = ["ffmpeg", "afconvert"]
        elif ext in all_images or f".{target}" in all_images:
            tool_list = ["sips", "magick", "pillow"]
        elif ext in all_docs or f".{target}" in all_docs:
            tool_list = ["textutil", "pandoc", "py_markdown", "py_pdf", "py_docx", "py_epub"]
        else:
            console.print(f"[red]Error: Don't know how to convert {ext} -> .{target}[/red]")
            raise typer.Exit(1)

    # Filter candidates by capability (does tool support output format?)
    capable_tools = [t for t in tool_list if _tool_supports_target(t, target)]
    
    if not capable_tools:
        # Fallback for PDF documents if textutil was suggested but can't do it
        if f".{target}" in EXTENSION_CATEGORIES["Documents"] and target == "pdf":
             capable_tools = ["pandoc", "py_pdf"]
        else:
             console.print(f"[red]Error: No known tool can convert {ext} -> .{target}[/red]")
             raise typer.Exit(1)

    tool = _resolve_tool(capable_tools)
    if tool is None:
        needed = [t for t in capable_tools if t in TOOLS_URLS]
        if needed:
            console.print(f"[NEED_INSTALL] {needed[0]}")
            console.print(f"[yellow]Missing tool: {needed[0]}. You can install it via the UI.[/yellow]")
        else:
            console.print(f"[red]Error: No available tool for {ext} -> .{target}[/red]")
        raise typer.Exit(1)

    console.print(f"[blue]Converting:[/blue] {Path(input_file).name} -> .{target} [dim](using {tool})[/dim]")

    success = False
    if tool == "python":
        ok = _convert_data(input_file, ext, target, output_path)
        if not ok: raise typer.Exit(1)
        success = True
    elif tool == "sips":
        sips_fmt = SIPS_FORMAT_ALIAS.get(target)
        if not sips_fmt:
            console.print(f"[red]Error: sips output .{target} not supported[/red]")
            raise typer.Exit(1)
        execute_shell_command(f'sips -s format {sips_fmt} "{input_file}" --out "{output_path}"')
        success = True
    elif tool == "afconvert":
        af_fmt = AFCONVERT_FORMATS_MAP.get(target)
        if not af_fmt:
            console.print(f"[red]Error: afconvert output .{target} not supported[/red]")
            raise typer.Exit(1)
        cmd = f'afconvert -f {af_fmt} -d {af_fmt.strip()} "{input_file}" "{output_path}"'
        if target in ("aac", "m4a"):
            cmd = f'afconvert -f m4af -d aac "{input_file}" "{output_path}"'
        execute_shell_command(cmd)
        success = True
    elif tool == "textutil":
        tu_fmt = target if target in TEXTUTIL_FORMATS else None
        if not tu_fmt:
            console.print(f"[red]Error: textutil output .{target} not supported[/red]")
            raise typer.Exit(1)
        execute_shell_command(f'textutil -convert {tu_fmt} -output "{output_path}" "{input_file}"')
        success = True
    elif tool == "ffmpeg":
        execute_shell_command(f'ffmpeg -y -i "{input_file}" "{output_path}"')
        success = True
    elif tool == "pandoc":
        execute_shell_command(f'pandoc "{input_file}" -o "{output_path}"')
        success = True
    elif tool == "magick":
        magick_cmd = "magick" if check_command_availability("magick") else "convert"
        execute_shell_command(f'{magick_cmd} "{input_file}" "{output_path}"')
        success = True
    # --- Bundled Python library converters ---
    elif tool == "pillow":
        ok = _convert_with_pillow(input_file, target, output_path)
        if not ok: raise typer.Exit(1)
        success = True
    elif tool == "pydub":
        ok = _convert_with_pydub(input_file, target, output_path)
        if not ok: raise typer.Exit(1)
        success = True
    elif tool in ("py_docx", "py_markdown", "py_pdf", "py_epub"):
        ok = _convert_document_python(input_file, ext, target, output_path)
        if not ok: raise typer.Exit(1)
        success = True
    elif tool in ("py_yaml", "py_xlsx", "py_toml"):
        ok = _convert_data(input_file, ext, target, output_path)
        if not ok: raise typer.Exit(1)
        success = True

    if success and os.path.exists(output_path):
        size = os.path.getsize(output_path)
        if size > 0:
            console.print(f"[green]Output:[/green] {output_path} ({_human_size(size)})")
        else:
            console.print(f"[yellow]Warning: Output file is empty.[/yellow]")


@app.command()
def organize(
    path: str,
    instructions: str = typer.Option("", "--instructions", "-i"),
    confirm: bool = typer.Option(False, "--confirm"),
):
    """Organize files in a folder."""
    if not os.path.isdir(path):
        console.print(f"[red]Error: Directory not found: {path}[/red]")
        raise typer.Exit(1)

    all_files = [f for f in os.listdir(path) if not f.startswith('.') and os.path.isfile(os.path.join(path, f))]
    if not all_files:
        console.print("[yellow]Folder is empty. Nothing to organize.[/yellow]")
        return

    if instructions.strip():
        plan = _ai_organize_plan(path, all_files, instructions)
    else:
        plan = {}
        for f in all_files:
            cat = _categorize_file(f)
            if cat not in plan: plan[cat] = []
            plan[cat].append(f)

    if not plan:
        console.print("[yellow]No files to move.[/yellow]")
        return

    total_moves = sum(len(files) for files in plan.values())
    console.print(f"\n[bold]Plan:[/bold] Move {total_moves} file(s) into {len(plan)} folder(s):\n")
    for folder, files in sorted(plan.items()):
        console.print(f"  [blue]{folder}/[/blue]")
        for f in files[:5]: console.print(f"    {f}")
        if len(files) > 5: console.print(f"    [dim]... and {len(files)-5} more[/dim]")

    if not confirm:
        console.print("\n[yellow]This is a preview. Re-run with --confirm to execute.[/yellow]")
        return

    console.print("")
    moved = 0
    for folder, files in plan.items():
        dest_dir = os.path.join(path, folder)
        os.makedirs(dest_dir, exist_ok=True)
        for f in files:
            src, dst = os.path.join(path, f), os.path.join(dest_dir, f)
            if os.path.exists(src) and not os.path.exists(dst):
                shutil.move(src, dst)
                moved += 1
    console.print(f"[green]Done. Moved {moved} file(s).[/green]")


def _ai_organize_plan(path: str, files: list, instructions: str) -> dict:
    # Detect image files for vision analysis
    image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}
    image_files = [f for f in files if Path(f).suffix.lower() in image_extensions]
    image_paths = [os.path.join(path, f) for f in image_files if os.path.exists(os.path.join(path, f))]
    
    if instructions:
        # User provided specific instructions - ONLY follow those
        if image_paths:
            prompt = f"""SYSTEM: You are a file organizer with vision capabilities. Output ONLY valid JSON.

FOLDER: "{path}"
FILES: {files[:100]}
USER INSTRUCTIONS: {instructions}

I'm showing you the actual images so you can see their CONTENT (not just filenames).
Analyze what's actually IN each image to make better organization decisions.

CRITICAL RULES:
1. Follow ONLY the user's specific instructions above
2. Use the image content (what you SEE in the images) to make decisions
3. Only move files that match the user's request based on their content
4. Leave all other files untouched (do NOT include them in the JSON)
5. If the user mentions a date, also look at image metadata/content for temporal clues
6. Return only the files requested, nothing else

OUTPUT FORMAT: JSON where keys are folder names and values are lists of filenames.
Example: {{"vacation_photos": ["IMG_1234.jpg", "IMG_5678.jpg"]}}
"""
        else:
            prompt = f"""SYSTEM: You are a file organizer. Output ONLY valid JSON.
FOLDER: "{path}"
FILES: {files[:100]}
USER INSTRUCTIONS: {instructions}

CRITICAL RULES:
1. Follow ONLY the user's specific instructions above
2. Only move files that match the user's request
3. Leave all other files untouched (do NOT include them in the JSON)
4. If the user mentions a date like "nov1st" or "nov 1", look for files from late October or early November
5. Return only the files requested, nothing else

OUTPUT FORMAT: JSON where keys are folder names and values are lists of filenames.
Example: {{"folder_name": ["file1.jpg", "file2.jpeg"]}}
"""
    else:
        # No specific instructions - organize everything by category
        if image_paths:
            prompt = f"""SYSTEM: You are a file organizer with vision capabilities. Output ONLY valid JSON.

FOLDER: "{path}"
FILES: {files[:100]}

I'm showing you the actual images so you can analyze their CONTENT.
Look at what's in each image and organize them into meaningful categories based on what you SEE.

TASK: Organize all files into logical category-based subfolders.
- For images: Use content-based categories (e.g., "Nature", "People", "Screenshots", "Memes", etc.)
- For other files: Use type-based categories (Documents, Videos, Music, etc.)

OUTPUT FORMAT: JSON where keys are folder names and values are lists of filenames.
Example: {{"Nature_Photos": ["sunset.jpg"], "Screenshots": ["screen1.png"], "Documents": ["report.pdf"]}}
"""
        else:
            prompt = f"""SYSTEM: You are a file organizer. Output ONLY valid JSON.
FOLDER: "{path}"
FILES: {files[:100]}
TASK: Organize all files into logical category-based subfolders (Images, Documents, Videos, etc.).
OUTPUT FORMAT: JSON where keys are folder names and values are lists of filenames.
"""
    
    # Use vision if we have images (limit to 10 for performance)
    response = get_llm_response(prompt, image_paths=image_paths[:10] if image_paths else None)
    try:
        cleaned = response.strip().replace("```json", "").replace("```", "")
        plan = json.loads(cleaned)
        if isinstance(plan, dict): return plan
    except (json.JSONDecodeError, ValueError, KeyError) as e:
        logging.warning(f"Failed to parse AI organize plan: {e}")
    return {}


@app.command()
def summarize(paths: List[str]):
    """Summarize files using AI."""
    if not paths:
        console.print("[red]Error: No files.[/red]")
        raise typer.Exit(1)
    
    contents = []
    for p in paths:
        if not os.path.isfile(p): continue
        try:
            with open(p, 'r', encoding='utf-8', errors='replace') as f:
                contents.append(f"FILE: {Path(p).name}\n{f.read(4000)}")
        except (OSError, IOError) as e:
            logging.warning(f"Failed to read file {p}: {e}")
        
    if not contents:
        console.print("[red]No readable files.[/red]")
        return

    console.print(f"[blue]Summarizing {len(contents)} file(s)...[/blue]")
    prompt = f"Summarize these files in 3 bullet points:\n\n{chr(10).join(contents)}"
    console.print(f"\n{get_llm_response(prompt)}")


@app.command()
def rename(
    paths: List[str],
    confirm: bool = typer.Option(False, "--confirm"),
):
    """Rename files using AI suggestions."""
    files = [p for p in paths if os.path.isfile(p)]
    if not files:
        console.print("[red]Error: No valid files.[/red]")
        raise typer.Exit(1)
    
    filenames = [Path(f).name for f in files[:30]]
    
    # Detect image files for vision analysis
    image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}
    image_files = [f for f in files if Path(f).suffix.lower() in image_extensions]
    
    if image_files:
        # Use vision model to analyze images
        prompt = f"""SYSTEM: You are a file renaming assistant with vision capabilities.

Analyze the provided images and suggest descriptive, clean filenames based on their CONTENT.
- Look at what's actually IN the image (objects, scenes, people, text, etc.)
- Create meaningful names that describe the image content
- Keep original file extensions
- Use underscores or hyphens for spaces
- Keep names concise (2-4 words maximum)

Current filenames: {filenames}

OUTPUT: Valid JSON map with "old_filename": "new_filename" pairs.
Example: {{"IMG_1234.jpg": "sunset_beach.jpg", "photo.png": "golden_retriever.png"}}
"""
        console.print(f"[blue]Analyzing {len(image_files)} image(s) with vision...[/blue]")
        response = get_llm_response(prompt, image_paths=image_files[:5])  # Limit to 5 images
    else:
        # Text-only prompt for non-image files
        prompt = f"""SYSTEM: File renamer. Output valid JSON map "old_name": "new_name".
FILES: {filenames}
TASK: Suggest clean names. Keep extensions.
"""
        console.print(f"[blue]Analyzing {len(files)} file(s)...[/blue]")
        response = get_llm_response(prompt)
    
    try:
        cleaned = response.strip().replace("```json", "").replace("```", "")
        renames = json.loads(cleaned)
    except (json.JSONDecodeError, ValueError) as e:
        console.print(f"[red]Error: AI failed to plan renames: {e}[/red]")
        return

    plan = []
    for p in files:
        old = Path(p).name
        if old in renames and renames[old] != old:
            plan.append((p, str(Path(p).parent / renames[old]), old, renames[old]))
            
    if not plan:
        console.print("[green]No renames needed.[/green]")
        return

    console.print("\n[bold]Plan:[/bold]")
    for _, _, old, new in plan:
        console.print(f"  {old} -> [green]{new}[/green]")

    if not confirm:
        console.print("\n[yellow]Preview only. Use --confirm to execute.[/yellow]")
        return

    count = 0
    for old_path, new_path, _, _ in plan:
        if not os.path.exists(new_path):
            os.rename(old_path, new_path)
            count += 1
    console.print(f"\n[green]Renamed {count} files.[/green]")


@app.command()
def ask(
    query: str = typer.Argument(..., help="Query"),
    paths: List[str] = typer.Argument(None),
):
    context = f"\nFiles: {paths}" if paths else ""
    prompt = f"Write code for: {query}{context}\nOutput ONLY ```python or ```bash block."
    response = get_llm_response(prompt)
    
    lang, code = extract_code_block(response)
    if lang == "python": execute_python_script(code)
    elif lang == "bash": execute_shell_command(code)
    else: console.print(f"[yellow]{response}[/yellow]")


# ============================================================
# AGENT-BASED DISPATCH / EXECUTE  (LLM-powered routing)
# ============================================================

AGENTS_DIR = Path(__file__).parent / "agents"


def _find_agents_dir() -> Path:
    """Locate the agents/ directory, checking multiple locations for dev and bundle contexts."""
    # 1. Environment override (e.g., set by Swift host)
    env_dir = os.environ.get("ALFRED_AGENTS_DIR")
    if env_dir:
        p = Path(env_dir)
        if p.is_dir():
            return p

    # 2. Sibling of alfred.py (works in both dev: cli/agents/ and bundle: Resources/agents/)
    sibling = Path(__file__).parent / "agents"
    if sibling.is_dir():
        return sibling

    # 3. Fallback: parent-of-parent (e.g., alfred.py in a subdirectory)
    parent_parent = Path(__file__).parent.parent / "agents"
    if parent_parent.is_dir():
        return parent_parent

    # 4. Last resort: return the default and let callers handle missing files
    return AGENTS_DIR


def _load_agent_prompt(agent_name: str) -> str:
    """Load the system prompt from agents/<name>.md"""
    agents_dir = _find_agents_dir()
    agent_file = agents_dir / f"{agent_name}.md"
    if not agent_file.exists():
        logging.warning(f"Agent prompt not found: {agent_file}")
        return f"You are Alfred's {agent_name} agent. Respond with valid JSON only."
    return agent_file.read_text(encoding="utf-8")


def _build_dispatch_context(agent: str, query: str, paths: list[str]) -> str:
    """Build the full prompt: agent system prompt + user query + file context."""
    system_prompt = _load_agent_prompt(agent)

    # Build file context
    file_context = ""
    if paths:
        file_details = []
        for p in paths:
            p_path = Path(p)
            if p_path.is_dir():
                files_in_dir = [
                    f.name for f in sorted(p_path.iterdir())
                    if not f.name.startswith('.') and f.is_file()
                ]
                file_details.append(f"FOLDER: {p}\nFILES IN FOLDER: {files_in_dir[:100]}")
            elif p_path.is_file():
                size = p_path.stat().st_size
                file_details.append(f"FILE: {p} (ext: {p_path.suffix}, size: {_human_size(size)})")
            else:
                file_details.append(f"PATH: {p} (not found)")
        file_context = "\n".join(file_details)

    user_message = f"""{system_prompt}

--- USER REQUEST ---
Query: {query}
{file_context}

Respond with ONLY valid JSON. No markdown fences, no explanation outside the JSON."""

    return user_message


@app.command()
def dispatch(
    agent: str = typer.Argument(..., help="Agent name: convert, organize, summarize, rename, command"),
    query: str = typer.Argument(..., help="Natural language query from the user"),
    paths: List[str] = typer.Argument(None, help="File or folder paths"),
):
    """Send a natural language query to an LLM agent. Returns a JSON plan for the UI to display."""
    valid_agents = {"convert", "organize", "summarize", "rename", "command"}
    if agent not in valid_agents:
        console.print(f'{{"action":"none","explanation":"Unknown agent: {agent}. Valid: {", ".join(valid_agents)}"}}')
        raise typer.Exit(1)

    file_list = list(paths) if paths else []
    prompt = _build_dispatch_context(agent, query, file_list)

    logging.info(f"Dispatch: agent={agent}, query={query!r}, paths={file_list}")

    # For organize agent with a folder, detect images for vision
    image_paths: list[str] | None = None
    if agent == "organize" and file_list:
        folder = Path(file_list[0])
        if folder.is_dir():
            image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}
            image_paths = [
                str(f) for f in folder.iterdir()
                if f.suffix.lower() in image_extensions and f.is_file()
            ][:10]  # limit

    # For rename agent, detect image files for vision
    if agent == "rename" and file_list:
        image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}
        image_paths = [p for p in file_list if Path(p).suffix.lower() in image_extensions][:5]

    response = get_llm_response(prompt, image_paths=image_paths if image_paths else None)

    # Clean up response — strip markdown fences if LLM added them
    cleaned = response.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("\n", 1)[-1] if "\n" in cleaned else cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned.rsplit("```", 1)[0]
    cleaned = cleaned.strip()

    # Validate it's JSON
    try:
        parsed = json.loads(cleaned)
        # Print clean JSON for the Swift UI to parse (use print, not console.print,
        # to avoid Rich wrapping/formatting — this output is machine-readable)
        print(json.dumps(parsed, ensure_ascii=False))
    except json.JSONDecodeError:
        # LLM didn't return valid JSON — wrap it
        fallback = {"action": "none", "explanation": f"Agent could not parse request: {cleaned[:200]}"}
        print(json.dumps(fallback, ensure_ascii=False))


@app.command()
def execute(
    plan_json: str = typer.Argument(..., help="JSON plan from dispatch"),
):
    """Execute a confirmed plan from dispatch. The plan JSON is the output of the dispatch command."""
    try:
        plan = json.loads(plan_json)
    except json.JSONDecodeError:
        console.print("[red]Error: Invalid JSON plan.[/red]")
        raise typer.Exit(1)

    action = plan.get("action", "none")

    if action == "none":
        console.print(f"[yellow]{plan.get('explanation', 'Nothing to do.')}[/yellow]")
        return

    # --- Convert ---
    if action == "convert":
        input_file = plan.get("input_file", "")
        target_format = plan.get("target_format", "")
        if not input_file or not target_format:
            console.print("[red]Error: Plan missing input_file or target_format.[/red]")
            raise typer.Exit(1)
        # Delegate to the existing convert command
        convert(input_file, target_format)
        return

    # --- Resize / compress image ---
    if action == "resize":
        input_file = plan.get("input_file", "")
        scale = float(plan.get("scale", 0.5))
        quality = int(plan.get("quality", 75))
        if not input_file:
            console.print("[red]Error: Plan missing input_file.[/red]")
            raise typer.Exit(1)
        src = Path(input_file)
        if not src.exists():
            console.print(f"[red]Error: File not found: {input_file}[/red]")
            raise typer.Exit(1)
        try:
            from PIL import Image as PILImage
        except ImportError:
            console.print("[red]Error: Pillow not available for resize.[/red]")
            raise typer.Exit(1)
        console.print(f"[blue]Resizing {src.name} (scale={scale}, quality={quality})...[/blue]")
        img = PILImage.open(src)
        if scale != 1.0:
            new_w = max(1, int(img.width * scale))
            new_h = max(1, int(img.height * scale))
            img = img.resize((new_w, new_h), PILImage.LANCZOS)
        # Save back to same file (or use a _small suffix to preserve original)
        stem = src.stem
        suffix = src.suffix.lower() or ".jpg"
        out = src.parent / f"{stem}_small{suffix}"
        save_kwargs: dict = {}
        if suffix in (".jpg", ".jpeg"):
            save_kwargs["quality"] = quality
            save_kwargs["optimize"] = True
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
        elif suffix == ".webp":
            save_kwargs["quality"] = quality
        elif suffix == ".png":
            save_kwargs["optimize"] = True
        img.save(out, **save_kwargs)
        original_size = src.stat().st_size
        new_size = out.stat().st_size
        console.print(f"  {src.name} -> [green]{out.name}[/green]")
        console.print(f"  Size: {_human_size(original_size)} -> [green]{_human_size(new_size)}[/green]")
        console.print(f"[green]Done.[/green]")
        return

    # --- Organize ---
    if action == "organize":
        folder = plan.get("folder", "")
        move_plan = plan.get("plan", {})
        if not folder or not move_plan:
            console.print("[red]Error: Plan missing folder or move plan.[/red]")
            raise typer.Exit(1)
        moved = 0
        for dest_name, files in move_plan.items():
            dest_dir = os.path.join(folder, dest_name)
            os.makedirs(dest_dir, exist_ok=True)
            for f in files:
                src = os.path.join(folder, f)
                dst = os.path.join(dest_dir, f)
                if os.path.exists(src) and not os.path.exists(dst):
                    shutil.move(src, dst)
                    moved += 1
        console.print(f"[green]Done. Moved {moved} file(s).[/green]")
        return

    # --- Summarize ---
    if action == "summarize":
        files = plan.get("files", [])
        style = plan.get("style", "brief")
        if not files:
            console.print("[red]Error: No files to summarize.[/red]")
            raise typer.Exit(1)
        contents = []
        for p in files:
            if not os.path.isfile(p):
                continue
            try:
                with open(p, 'r', encoding='utf-8', errors='replace') as f:
                    contents.append(f"FILE: {Path(p).name}\n{f.read(4000)}")
            except (OSError, IOError):
                pass
        if not contents:
            console.print("[red]No readable files.[/red]")
            return
        style_instructions = {
            "brief": "Summarize in 3 concise bullet points per file.",
            "detailed": "Provide a detailed paragraph-level breakdown.",
            "comparison": "Compare and contrast the files, highlighting similarities and differences.",
            "explain": "Explain the content as if teaching someone. Cover what it does and why.",
        }
        instruction = style_instructions.get(style, style_instructions["brief"])
        prompt = f"{instruction}\n\n{chr(10).join(contents)}"
        console.print(f"[blue]Summarizing {len(contents)} file(s) ({style})...[/blue]")
        console.print(f"\n{get_llm_response(prompt)}")
        return

    # --- Rename ---
    if action == "rename":
        renames = plan.get("renames", {})
        if not renames:
            console.print("[green]No renames needed.[/green]")
            return
        # We need the original full paths — renames dict has filename: new_filename
        # The paths are not in the plan, so we infer from the keys
        count = 0
        for old_name, new_name in renames.items():
            # Try to find the file — could be absolute or relative
            old_path = Path(old_name)
            if not old_path.exists():
                # Maybe it's just a filename — skip if we can't find it
                console.print(f"[yellow]Skipped: {old_name} (not found)[/yellow]")
                continue
            new_path = old_path.parent / new_name
            if not new_path.exists():
                os.rename(str(old_path), str(new_path))
                console.print(f"  {old_path.name} -> [green]{new_name}[/green]")
                count += 1
        console.print(f"\n[green]Renamed {count} file(s).[/green]")
        return

    # --- Run (command agent) ---
    if action == "run":
        language = plan.get("language", "")
        code = plan.get("code", "")
        if not code:
            console.print("[red]Error: No code to execute.[/red]")
            raise typer.Exit(1)
        if language == "python":
            execute_python_script(code)
        elif language == "bash":
            execute_shell_command(code)
        else:
            console.print(f"[red]Error: Unknown language '{language}'[/red]")
        return

    console.print(f"[red]Error: Unknown action '{action}'[/red]")
    raise typer.Exit(1)


if __name__ == "__main__":
    _init_config()
    app()
